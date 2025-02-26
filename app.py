# -*- coding: utf-8 -*-
import random,os,json,io,re,zipfile,tempfile
import pandas as pd
import streamlit as st
import streamlit_toggle as tog
from langchain.llms import OpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
import openai 
import fitz
import docx
from gtts import gTTS
import PyPDF2
from PyPDF2 import PdfReader
from utils import text_to_docs
from langchain import PromptTemplate, LLMChain
from langchain.chat_models import ChatOpenAI
from langchain.chains import ConversationChain
from langchain.memory import ConversationBufferMemory
from langchain.memory import ConversationSummaryBufferMemory
from langchain.chains.conversation.prompt import ENTITY_MEMORY_CONVERSATION_TEMPLATE
from langchain.chains.conversation.memory import ConversationEntityMemory
from langchain.callbacks import get_openai_callback
from io import StringIO
from io import BytesIO
from usellm import Message, Options, UseLLM
#from playsound import playsound
#from langchain.text_splitter import CharacterTextSplitter
#from langchain.embeddings.openai import OpenAIEmbeddings
#from langchain.chains.summarize import load_summarize_chain
#import os
#import pyaudio
#import wave
#from langchain.document_loaders import UnstructuredPDFLoader
#import streamlit.components.v1 as components
#from st_custom_components import st_audiorec, text_to_docs
#import sounddevice as sd
#from scipy.io.wavfile import write

# Setting Env
if st.secrets["OPENAI_API_KEY"] is not None:
    os.environ["OPENAI_API_KEY"] = st.secrets["OPENAI_API_KEY"]
else:
    os.environ["OPENAI_API_KEY"] = os.environ.get("OPENAI_API_KEY")

@st.cache_data
def show_pdf(file_path):
    with open(file_path,"rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
        pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="1000px" type="application/pdf"></iframe>'
        st.markdown(pdf_display, unsafe_allow_html=True)

@st.cache_data
def pdf_to_bytes(pdf_file_):
    with open(pdf_file_,"rb") as pdf_file:
        pdf_content = pdf_file.read()
        pdf_bytes_io = io.BytesIO(pdf_content)
    return pdf_bytes_io

@st.cache_data
def read_pdf_files(path):
    pdf_files =[]
    directoty_path = path
    files = os.listdir(directoty_path)
    for file in files:
        if file.lower().endswith('.pdf'):
            pdf_files.append(file)
    return pdf_files


@st.cache_data
def merge_pdfs(pdf_list):
    """
    Helper function to merge PDFs
    """
    pdf_merger = PyPDF2.PdfMerger()
    for pdf in pdf_list:
        pdf_document = PyPDF2.PdfReader(pdf)
        pdf_merger.append(pdf_document)
    output_pdf = BytesIO()
    pdf_merger.write(output_pdf)
    pdf_merger.close()
    return output_pdf


@st.cache_data
def usellm(prompt):

    service = UseLLM(service_url="https://usellm.org/api/llm")
    messages = [
      Message(role="system", content="You are a fraud analyst, who is an expert at finding out suspicious activities"),
      Message(role="user", content=f"{prompt}"),
      ]
    options = Options(messages=messages)
    response = service.chat(options)
    return response.content

@st.cache_data
def process_text(text):
    # Add your custom text processing logic here
    processed_text = text
    return processed_text

@st.cache_resource
def embed(model_name):
    hf_embeddings = HuggingFaceEmbeddings(model_name=model_name)
    return hf_embeddings

@st.cache_data
def embedding_store(pdf_files):
    merged_pdf = merge_pdfs(pdf_files)
    final_pdf = PyPDF2.PdfReader(merged_pdf)
    text = ""
    for page in final_pdf.pages:
        text += page.extract_text()
    texts =  text_splitter.split_text(text)
    docs = text_to_docs(texts)
    docsearch = FAISS.from_documents(docs, hf_embeddings)
    return docs, docsearch
    
@st.cache_data
def merge_and_extract_text(pdf_list):
    """
    Helper function to merge PDFs and extract text
    """
    pdf_merger = PyPDF2.PdfMerger()
    for pdf in pdf_list:
        with open(pdf, 'rb') as file:
            pdf_merger.append(file)
    output_pdf = BytesIO()
    pdf_merger.write(output_pdf)
    pdf_merger.close()
    
    # Extract text from merged PDF
    merged_pdf = PyPDF2.PdfReader(output_pdf)
    all_text = []
    for page in merged_pdf.pages:
        text = page.extract_text()
        all_text.append(text)
    
    return ' '.join(all_text)

def reset_session_state():
    session_state = st.session_state
    session_state.clear()

# def merge_and_extract_text(pdf_list):
#     merged_pdf = fitz.open()
#     # Merge the PDF files
#     for pdf_file in pdf_list:
#         pdf_document = fitz.open(pdf_file)
#         merged_pdf.insert_pdf(pdf_document)
#     # Create an empty string to store the extracted text
#     merged_text = ""
#     # Extract text from each page of the merged PDF
#     for page_num in range(merged_pdf.page_count):
#         page = merged_pdf[page_num]
#         text = page.get_text()
#         merged_text += text
#     # Close the merged PDF
#     merged_pdf.close()
#     return merged_text


@st.cache_data
def render_pdf_as_images(pdf_file):
    """
    Helper function to render PDF pages as images
    """
    pdf_images = []
    pdf_document = fitz.open(stream=pdf_file.read(), filetype="pdf")
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        img = page.get_pixmap()
        img_bytes = img.tobytes()
        pdf_images.append(img_bytes)
    pdf_document.close()
    return pdf_images

# To check if pdf is searchable
def is_searchable_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            if page.extract_text():
                return True

    return False


# Function to add checkboxes to the DataFrame
@st.cache_data
def add_checkboxes_to_dataframe(df):
    # Create a new column 'Select' with checkboxes
    checkbox_values = [True] * (len(df) - 1) + [False]  # All True except the last row
    df['Select'] = checkbox_values
    return df

# convert scanned pdf to searchable pdf
def convert_scanned_pdf_to_searchable_pdf(input_file):
    """
     Convert a Scanned PDF to Searchable PDF

    """
    # Convert PDF to images
    print("Running OCR")
    images = convert_from_path(input_file)

    # Preprocess images using OpenCV
    for i, image in enumerate(images):
        # Convert image to grayscale
        image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2GRAY)

        # Apply thresholding to remove noise
        _, image = cv2.threshold(image, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

        # Enhance contrast
        image = cv2.equalizeHist(image)

        # Save preprocessed image
        cv2.imwrite(f'{i}.png', image)

    # Perform OCR on preprocessed images using Tesseract
    text = ''
    for i in range(len(images)):
        image = cv2.imread(f'{i}.png')
        text += pytesseract.image_to_string(image)
    
    return text


# Setting globals
if "visibility" not in st.session_state:
    st.session_state.visibility = "visible"
    st.session_state.disabled = True
if "stored_session" not in st.session_state:
    st.session_state["stored_session"] = []
if "tmp_table" not in st.session_state:
    st.session_state.tmp_table=pd.DataFrame()
if "tmp_summary" not in st.session_state:
    st.session_state["tmp_summary"] = ''
if "case_num" not in st.session_state:
    st.session_state.case_num = ''
if "fin_opt" not in st.session_state:
    st.session_state.fin_opt = ''
if "context_1" not in st.session_state:
    st.session_state.context_1 = ''

# reading files from local directory from fetch evidence button
directoty_path = "data/"
fetched_files = read_pdf_files(directoty_path)


# Apply CSS styling to resize the buttons
st.markdown("""
    <style>
        .stButton button {
            width: 145px;
            height: 35px;
        }
    </style>
""", unsafe_allow_html=True)

@st.cache_data
def add_footer_with_fixed_text(doc, footer_text):
    # Create a footer object
    footer = doc.sections[0].footer

    # Add a paragraph to the footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

    # Set the fixed text in the footer
    paragraph.text = footer_text

    # Add a page number field to the footer
    run = paragraph.add_run()
    fld_xml = f'<w:fldSimple {nsdecls("w")} w:instr="PAGE"/>'
    fld_simple = parse_xml(fld_xml)
    run._r.append(fld_simple)

    # Set the alignment of the footer text
    paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

@st.cache_data
def create_filled_box_with_text(color, text):
    box_html = f'<div style="flex: 1; height: 100px; background-color: {color}; display: flex; align-items: center; justify-content: center;">{text}</div>'
    st.markdown(box_html, unsafe_allow_html=True)

@st.cache_data
def create_zip_file(file_paths, zip_file_name):
    with zipfile.ZipFile(zip_file_name, 'w') as zipf:
        for file_path in file_paths:
            zipf.write(file_path, os.path.basename(file_path))

# Addding markdown styles(Global)
st.markdown("""
<style>
.big-font {
    font-size:60px !important;
}
</style>
""", unsafe_allow_html=True)


# Set Sidebar
st.markdown("""
<style>
    [data-testid=stSidebar] {
        background-color: FBFBFB;
    }
</style>
""", unsafe_allow_html=True)

st.title("Suspicious Activity Reporting Assistant")
with st.sidebar:
    # st.sidebar.write("This is :blue[test]")
    # Navbar
    st.markdown('<link rel="stylesheet" href="https://maxcdn.bootstrapcdn.com/bootstrap/4.0.0/css/bootstrap.min.css" integrity="sha384-Gn5384xqQ1aoWXA+058RXPxPg6fy4IWvTNh0E263XmFcJlSAwiGgFAW/dAiS6JXm" crossorigin="anonymous">', unsafe_allow_html=True)

    st.markdown("""
    <nav class="navbar fixed-top navbar-expand-lg navbar-dark" style="background-color: #000000;">
    <button class="navbar-toggler" type="button" data-toggle="collapse" data-target="#navbarNav" aria-controls="navbarNav" aria-expanded="false" aria-label="Toggle navigation">
        <span class="navbar-toggler-icon"></span>
    </button>
    <style>
    .navbar-brand img {
      max-height: 50px; /* Adjust the height of the logo */
      width: auto; /* Allow the width to adjust based on the height */
    }
    </style>
    <div class="collapse navbar-collapse" id="navbarNav">
        <ul class="navbar-nav">
        <li class="nav-item active">
            <a class="navbar-brand" href="#">
                <img src="https://www.exlservice.com/themes/exl_service/exl_logo_rgb_orange_pos_94.png" width="50" height="30" alt="">
                <span class="sr-only">(current)</span>
                <strong>| Operations Process Automation</strong>
            </a>
        </li>
        </ul>
    </div>
    </nav>
    """, unsafe_allow_html=True)

    st.markdown("""
    <nav class="navbar fixed-bottom navbar-expand-lg navbar-dark" style="background-color: #000000;">
    <button class="navbar-toggler" type="button" data-toggle="collapse" data-target="#navbarNav" aria-controls="navbarNav" aria-expanded="false" aria-label="Toggle navigation">
        <span class="navbar-toggler-icon"></span>
    </button>
    <div class="collapse navbar-collapse" id="navbarNav">
        <ul class="navbar-nav">
        <li class="nav-item active">
        <!--p style='color: white;'><b>Powered by EXL</b></p--!>
        <p style='color: white;'> <strong>Powered by EXL</strong> </p>
            <!--a class="nav-link disabled" href="#">
                <img src="https://www.exlservice.com/themes/exl_service/exl_logo_rgb_orange_pos_94.png" width="50" height="30" alt="">
                <span class="sr-only">(current)</span>
            </a--!>
        </li>
        </ul>
    </div>
    </nav>
    """, unsafe_allow_html=True)

    # Add the app name
    st.sidebar.markdown('<p class="big-font">SARA</p>', unsafe_allow_html=True)
    # st.sidebar.header("SARA")

    # Add a drop-down for case type
    options = ["Select Case Type", "Fraud transaction dispute", "AML"]
    selected_option_case_type = st.sidebar.selectbox("", options)
    st.markdown("---")
    
    # Add a single dropdown
    options = ["Select Case ID", "SAR-2023-24680", "SAR-2023-13579", "SAR-2023-97531", "SAR-2023-86420", "SAR-2023-24681"]
    selected_option = st.sidebar.selectbox("", options)
    # Add the image to the sidebar below options
    st.sidebar.image("MicrosoftTeams-image (7).png", use_column_width=True)

    
# Assing action to the main section
if selected_option_case_type == "Select Case Type":
    st.header("")
elif selected_option_case_type == "Fraud transaction dispute":
    st.markdown("### :blue[Fraud transaction dispute]")
elif selected_option_case_type == "AML":
    st.markdown("### :red[Anti-Money Laundering]")
st.markdown('---')

# Redirect to Merge PDFs page when "Merge PDFs" is selected
if selected_option == "SAR-2023-24680":
    st.session_state.case_num = "SAR-2023-24680"
    # st.header("Merge Documents")
    # st.write("Upload multiple document files and merge them into one doc.")

    # Upload PDF files
    # st.subheader("Upload Case Files")
    # st.markdown(f"**Case No: {st.session_state.case_num}**")
    # st.markdown("""
    #     | Case No.                  | Case Type                 | Customer Name             | Case Status             | Open Date              |
    #     | ------------------------  | ------------------------- | ------------------------- | ------------------------|------------------------|
    #     | SAR-2023-24680            | Fraud Transaction Dispute | John Brown                | In Progress             | 12/10/2020             |
    #     """)

    col1,col2 = st.columns(2)
    # Row 1
    with col1:
        st.markdown("**Case number&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** SAR-2023-24680")
        st.markdown("**Customer name  :** John Brown")


    with col2:
        st.markdown("**Case open date&nbsp;&nbsp;&nbsp;&nbsp;:** Feb 02, 2021")
        st.markdown("**Case type&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** Fraud transaction")


    # Row 2
    with col1:
        st.markdown("**Customer ID&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** 9659754")


    with col2:
        st.markdown("**Case Status&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** Open")


    # Evidence uploader/Fetch    
    st.header("Upload Evidence")
    # Showing files record
    # data_source = ["Customer details", "Transaction details", "Fraud dispute details", "Other Files"]
    # data_file_df = pd.DataFrame(data_source, columns=["File Name"])
    # data_file_df = data_file_df.reset_index(drop=True)

    # df_s_with_checkboxes = add_checkboxes_to_dataframe(data_file_df)

    # # Iterate through each row and add checkboxes
    # for index, row in df_s_with_checkboxes.iterrows():
    #     if index < -1:
    #         checkbox_state = st.checkbox(f" {row['File Name']}", value=True)
    #         df_s_with_checkboxes.loc[index, 'Select'] = checkbox_state
    #     else:
    #         st.checkbox(f"{row['File Name']}", value=False)


    if selected_option:
        # Create two columns
        col1_up, col2_up = st.tabs(["Fetch Evidence", "Upload Evidence"])
        with col1_up:
            # Set the color
            # st.markdown(
            #     """
            #     <div style="display: flex; justify-content: center; align-items: center; height: 48px; border: 1px solid #ccc; border-radius: 5px; background-color: #f2f2f2;">
            #         <span style="font-size: 16px;  ">Fetch Evidence</span>
            #     </div>
            #     """,
            #     unsafe_allow_html=True
            # )
            if 'clicked' not in st.session_state:
                st.session_state.clicked = False
            
            def set_clicked():
                st.session_state.clicked = True
                st.session_state.disabled = True
            
            st.button('Fetch Evidence', on_click=set_clicked)

            if st.session_state.clicked:
                # st.write("Evidence Files:") 
                # st.markdown(html_str, unsafe_allow_html=True)
                
                # Showing files
                # show_files = fetched_files.copy()
                # show_files = show_files + ['Other.pdf']
                # files_frame = pd.DataFrame(show_files, columns=["File Name"])
                # # files_frame["Select"] = [True for _ in range(len(files_frame))]
                # files_frame = files_frame.reset_index(drop=True)

                # # Add checkboxes to the DataFrame
                # df_with_checkboxes = add_checkboxes_to_dataframe(files_frame)
               
                # # Iterate through each row and add checkboxes
                # for index, row in df_with_checkboxes.iterrows():
                #     if index < len(df_with_checkboxes) - 1:
                #         checkbox_state = st.checkbox(f" {row['File Name']}", value=True)
                #         df_with_checkboxes.loc[index, 'Select'] = checkbox_state
                #     else:
                #         st.checkbox(f"{row['File Name']}", value=False)



                # st.dataframe(files_frame)
                # st.write(df_reset.to_html(index=False), unsafe_allow_html=True)
                # st.markdown(files_frame.style.hide(axis="index").to_html(), unsafe_allow_html=True)
                
                
                
                #select box to select file
                selected_file_name = st.selectbox(":blue[Select a file to View]",fetched_files)
                st.write("Selected File: ", selected_file_name)
                st.session_state.disabled = False
                if selected_file_name:
                    selected_file_path = os.path.join(directoty_path, selected_file_name)
                    #converting pdf data to bytes so that render_pdf_as_images could read it
                    file = pdf_to_bytes(selected_file_path)
                    pdf_images = render_pdf_as_images(file)
                    #showing content of the pdf
                    st.subheader(f"Contents of {selected_file_name}")
                    for img_bytes in pdf_images:
                        st.image(img_bytes, use_column_width=True)



        with col2_up:
            pdf_files = st.file_uploader("", type=["pdf","png","jpeg","docx","xlsx"], accept_multiple_files=True)
            


            # Show uploaded files in a dropdown
            if pdf_files:
                st.subheader("Uploaded Files")
                file_names = [file.name for file in pdf_files]
                selected_file = st.selectbox(":blue[Select a file]", file_names)
                # Enabling the button
                st.session_state.disabled = False
                # Display selected PDF contents
                if selected_file:
                    selected_pdf = [pdf for pdf in pdf_files if pdf.name == selected_file][0]
                    pdf_images = render_pdf_as_images(selected_pdf)
                    st.subheader(f"Contents of {selected_file}")
                    for img_bytes in pdf_images:
                        st.image(img_bytes, use_column_width=True)

    tmp_dir_ = tempfile.mkdtemp()
    temp_file_path= []

    for uploaded_file in pdf_files:
        file_pth = os.path.join(tmp_dir_, uploaded_file.name)
        with open(file_pth, "wb") as file_opn:
            file_opn.write(uploaded_file.getbuffer())
            temp_file_path.append(file_pth)


    for fetched_pdf in fetched_files:
        file_pth = os.path.join('data/', fetched_pdf)
        # st.write(file_pth)
        temp_file_path.append(file_pth) 

    #combining files in fetch evidence and upload evidence
    if temp_file_path:
        if pdf_files and fetched_files:
            file_names = [file.name for file in pdf_files]
            file_names = file_names + fetched_files
            pdf_files_ = file_names
        elif fetched_files:
            pdf_files_ = fetched_files
        elif pdf_files:
            pdf_files_ = pdf_files
        else: pass

    model_name = "sentence-transformers/all-MiniLM-L6-v2"
    # model_name = "hkunlp/instructor-large"
    
    # Memory setup
    llm = ChatOpenAI(temperature=0.1)
    memory = ConversationSummaryBufferMemory(llm=llm, max_token_limit=500)
    conversation = ConversationChain(llm=llm, memory =memory,verbose=False)
    
    
    # Adding condition on embedding
    try:
        if temp_file_path:
            hf_embeddings = embed(model_name) 
        else:
            pass
    except NameError:
        pass
    
    # Chunking with overlap
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size = 1000,
        chunk_overlap  = 100,
        length_function = len,
        separators=["\n\n", "\n", " ", ""]
    )
    #text_splitter = CharacterTextSplitter.from_tiktoken_encoder(chunk_size=100, chunk_overlap=0)
    #texts = ''
    
    # @st.cache_data
    # def embedding_store(file):
    #     # save file
    #     pdf_reader = PdfReader(file)
    #     text = ""
    #     for page in pdf_reader.pages:
    #         text += page.extract_text()
    #     #st.write(text)
    #     texts =  text_splitter.split_text(text)
    #     docs = text_to_docs(texts)
    #     #st.write(texts)
    #     docsearch = FAISS.from_documents(docs, hf_embeddings)
    #     return docs, docsearch
    
    

else:
    # Disabling the button
    st.session_state.disabled = True
    st.session_state.case_num = selected_option    

# Creating header
col1,col2 = st.columns(2)
with col1:
    st.subheader('Pre-Set Questionnaire')
    # Create a Pandas DataFrame with your data

    data = {'Questions': [" What is the Victim's Name?",' Has any suspect been reported?',' List the Merchant name',' How was the bank notified?',' When was the bank notified?',' What is the Fraud Type?',' When did the fraud occur?',' Was the disputed amount greater than 5000 USD?',' What type of cards are involved?',' Was the police report filed?']}
    df_fixed = pd.DataFrame(data)
    df_fixed.index = df_fixed.index +1
with col2:
    # Create a checkbox to show/hide the table
    cols1, cols2, cols3, cols4 = st.columns([1,1,1,1])
    with cols1:
        show_table = tog.st_toggle_switch(label="", 
                            key="Key1", 
                            default_value=False, 
                            label_after = False, 
                            inactive_color = '#D3D3D3', 
                            active_color="#11567f", 
                            track_color="#29B5E8"
                            )
    # Show the table if the checkbox is ticked
    if show_table:
        # st.write(df_fixed)
        # st.dataframe(df_fixed, width=1000)
        df_fixed["S.No."] = df_fixed.index
        df_fixed = df_fixed.loc[:,['S.No.','Questions']]
        st.markdown(df_fixed.style.hide(axis="index").to_html(), unsafe_allow_html=True)

with st.spinner('Wait for it...'):
    if st.button("Generate Insights",disabled=st.session_state.disabled):
        if pdf_files is not None:
            # File handling logic
            _, docsearch = embedding_store(temp_file_path)
            queries ="Please provide the following information regarding the possible fraud case: What is the name of the customer name,\
            has any suspect been reported, list the merchant name, how was the bank notified, when was the bank notified, what is the fraud type,\
            when did the fraud occur, was the disputed amount greater than 5000 USD, what type of cards are involved, was the police report filed,\
            and based on the evidence, is this a suspicious activity(Summarize all the questions asked prior to this in a detailed manner),that's the answer of\
            whether this is a suspicious activity\
            "
            contexts = docsearch.similarity_search(queries, k=5) 
            prompts = f" Give a the answer to the below questions as truthfully and in as detailed in the form of sentences\
            as possible as per given context only,\n\n\
                    1. What is the Victim's Name?\n\
                    2. Has any suspect been reported?\n\
                    3. List the Merchant name\n\
                    4. How was the bank notified?\n\
                    5. When was the bank notified?\n\
                    6. What is the Fraud Type?\n\
                    7. When did the fraud occur?\n\
                    8. Was the disputed amount greater than 5000 USD?\n\
                    9. What type of cards are involved?\n\
                    10. Was the police report filed?\n\
                Context: {contexts}\n\
                Response (in the python dictionary format\
                where the dictionary key would carry the questions and its value would have a descriptive answer to the questions asked): "
                

            response = usellm(prompts)
            # st.write(response)
            # memory.save_context({"input": f"{queries}"}, {"output": f"{response}"})
            # st.write(response)
            # st.write(memory.load_memory_variables({}))



            # Convert the response in dictionary from tbl
            # prompt_conv = f" Convert the tabular data into a python dictionary\
            #     context: {response}\
            #     Response (give me the response in the form of a python dictionary with questions exactly as it is): "
            # resp_dict = usellm(prompt_conv)
            # st.write(response)
            resp_dict_obj = json.loads(response)
            res_df = pd.DataFrame(resp_dict_obj.items(), columns=['Question','Answer'])
            # st.table(res_df)
            try:
                res_df.Question = res_df.Question.apply(lambda x: x.split(".")[1])
                res_df.index = res_df.index + 1
                df_base = res_df.copy(deep=True)
                df_base["S.No."] = df_base.index
                df_base = df_base.loc[:,['S.No.','Question','Answer']]
            except IndexError:
                pass
            # st.table(res_df)
            st.markdown(df_base.style.hide(axis="index").to_html(), unsafe_allow_html=True)
            # st.write(resp_dict_obj)
            st.session_state["tmp_table"] = pd.concat([st.session_state.tmp_table, res_df], ignore_index=True)
st.markdown("---")

# For input box outside of template4
try:
    if temp_file_path:
        docs, docsearch = embedding_store(temp_file_path)
    else:
        pass
except Exception:
    pass


# Text Input
st.subheader("Ask Additional Questions")
query = st.text_input(':blue[Please ask below the additional case questions.]',disabled=st.session_state.disabled)
text_dict = {}
@st.cache_data
def LLM_Response():
    llm_chain = LLMChain(prompt=prompt, llm=llm)
    response = llm_chain.run({"query":query, "context":context})
    return response

with st.spinner('Getting you information...'):      
    if query:
        # Text input handling logic
        #st.write("Text Input:")
        #st.write(text_input)

        context_1 = docsearch.similarity_search(query, k=5)
        st.session_state.context_1 = context_1
        if query.lower() == "what is the victim's name?":
            prompt_1 = f'''Perform Name Enitity Recognition to identify the Customer name as accurately as possible, given the context. The Customer can also be referenced as the Victim or the person with whom the Fraud has taken place.\n\n\
                        Question: {query}\n\
                        Context: {context_1}\n\
                        Response: '''

            
        elif query.lower() == "what is the suspect's name?":
            prompt_1 = f'''Perform Name Enitity Recognition to identify the Suspect name as accurately as possible, given the context. Suspect is the Person who has committed the fraud with the Customer. Respond saying "The Suspect Name is not Present" if there is no suspect in the given context.\n\n\
                        Question: {query}\n\
                        Context: {context_1}\n\
                        Response: '''

            
        elif query.lower() == "list the merchant name":
            prompt_1 = f'''Perform Name Enitity Recognition to identify all the Merchant Organizations as accurately as possible, given the context. A merchant is a type of business or organization that accepts payments from the customer account. Give a relevant and concise response.\n\n\
                        Question: {query}\n\
                        Context: {context_1}\n\
                        Response: '''

            
        elif query.lower() == "how was the bank notified?":
            prompt_1 = f''' You need to act as a Financial analyst to identify how was the bank notified of the Supicious or Fraud event with in the given context. The means of communication can be a call, an email or in person. Give a relevant and concise response.\n\n\
                        Question: {query}\n\
                        Context: {context_1}\n\
                        Response: '''

            
        elif query.lower() == "when was the bank notified?":
            prompt_1 = f''' You need to act as a Financial analyst to identify the when the bank was notified of the Fraud i.e., the disputed date. Given the context, provide a relevant and concise response.\n\n\
                        Question: {query}\n\
                        Context: {context_1}\n\
                        Response: '''

            
        elif query.lower() == "what type of fraud is taking place?":
            prompt_1 = f''' You need to act as a Financial analyst to identify the type of fraud or suspicious activity has taken place amd summarize it, within the given context. Also mention the exact fraud code. Give a relevant and concise response.\n\n\
                        Question: {query}\n\
                        Context: {context_1}\n\
                        Response: '''

        
        elif query.lower() == "when did the fraud occur?":
            prompt_1 = f''' You need to act as a Financial analyst to identify the when the did the fraud occur i.e., the Transaction Date. Given the context, provide a relevant and concise response.\n\n\
                        Question: {query}\n\
                        Context: {context_1}\n\
                        Response: '''

                
        elif query.lower() == "was the disputed amount greater than 5000 usd?":
            prompt_1 = f''' You need to act as a Financial analyst to identify the disputed amount and perform a mathematical calculation to check if the disputed amount is greater than 5000 or no, given the context. Give a relevant and concise response.\n\n\
                        Question: {query}\n\
                        Context: {context_1}\n\
                        Response: '''

            
        elif query.lower() == "what type of cards are involved?":
            prompt_1 = f''' You need to act as a Financial analyst to identify the type of card and card's brand involved, given the context. On a higher level the card can be a Credit or Debit Card. VISA, MasterCard or American Express, Citi Group, etc. are the different brands with respect to a Credit card or Debit Card . Give a relevant and concise response.\n\n\
                        Question: {query}\n\
                        Context: {context_1}\n\
                        Response: '''

            
        elif query.lower() == "was the police report filed?":
            prompt_1 = f''' You need to act as a Financial analyst to identify if the police was reported of the Fraud activity, given the context. Give a relevant and concise response.\n\n\
                        Question: {query}\n\
                        Context: {context_1}\n\
                        Response: '''

                
        elif query.lower() == "Is this a valid SAR case?":
            prompt_1 = f''' You need to act as a Financial analyst to check if this is a SAR or not, given the following context, if the transaction amount is less than 5000 USD we cannot categorize this as SAR (Suspicious activity Report).Give a relevant and concise response. \n\n\
                        Question: {query}\n\
                        Context: {context_1}\n\
                        Response: '''

            
        else:
            prompt_1 = f'''Act as a financial analyst and give concise answer to below Question as truthfully as possible, with given Context.\n\n\
                        Question: {query}\n\
                        Context: {context_1}\n\                      
                        Response: '''


        #prompt = PromptTemplate(template=prompt, input_variables=["query", "context"])
        response = usellm(prompt_1) #LLM_Response()
        text_dict[query] = response
        # resp_dict_obj.update(text_dict)
        st.write(response)
        if response:
            df = pd.DataFrame(text_dict.items(), columns=['Question','Answer'])
        else:
            df = pd.DataFrame()
        # st.session_state.tmp_table = pd.concat([tmp_table, tmp_table], ignore_index=True)
        # st.write(text_dict.items())
        st.session_state["tmp_table"] = pd.concat([st.session_state.tmp_table, df], ignore_index=True)
        st.session_state.tmp_table.drop_duplicates(subset=['Question'])
        # st.table(st.session_state.tmp_table)


with st.spinner('Summarization ...'):  
    if st.button("Summarize",disabled=st.session_state.disabled):
        summ_dict = st.session_state.tmp_table.set_index('Question')['Answer'].to_dict()
        # chat_history = resp_dict_obj['Summary']
        memory = ConversationSummaryBufferMemory(llm=llm, max_token_limit=300)
        memory.save_context({"input": "This is the entire chat summary"}, 
                        {"output": f"{summ_dict}"})
        conversation = ConversationChain(
        llm=llm, 
        memory = memory,
        verbose=True)
        st.session_state["tmp_summary"] = conversation.predict(input="Give me a detailed summary of the above texts in a single paragraph without anything additional other than the overall content. Please don't include words like these: 'chat summary', 'includes information' in my final summary.")
        # showing the text in a textbox
        # usr_review = st.text_area("", value=st.session_state["tmp_summary"])
        # if st.button("Update Summary"):
        #     st.session_state["fin_opt"] = usr_review
        st.write(st.session_state["tmp_summary"])


with st.spinner("Downloading...."):
# if st.button("Download Response", disabled=st.session_state.disabled):
    # Create a Word document with the table and some text
    if st.session_state["tmp_summary"]:
        st.session_state.disabled=False
        
    try:
        # initiate the doc file
        doc = docx.Document()
        # doc.add_section(WD_SECTION.NEW_PAGE)
        doc.add_heading(f"Case No.: {st.session_state.case_num}",0)

        # Add a subheader for case details
        subheader_case = doc.add_paragraph("Case Details")
        subheader_case.style = "Heading 2"
        # Addition of case details
        paragraph = doc.add_paragraph(" ")
        case_info = {
            "Case Number                            ": " SAR-2023-24680",
            "Customer Name                       ": " John Brown",
            "Customer ID                              ": " 9659754",
            "Case open date                         ": " Feb 02, 2021",
            "Case Type                                  ": " Fraud Transaction",
            "Case Status                                ": " Open"
        }
        for key_c, value_c in case_info.items():
            doc.add_paragraph(f"{key_c}: {value_c}")
        paragraph = doc.add_paragraph(" ")

        # Add a subheader for customer info to the document ->>
        subheader_paragraph = doc.add_paragraph("Customer Information")
        subheader_paragraph.style = "Heading 2"
        paragraph = doc.add_paragraph(" ")

        # Add the customer information
        customer_info = {
            "Name                                           ": " John Brown",
            "Address                                      ": " 858 3rd Ave, Chula Vista, California, 91911 US",
            "Phone                                          ": " (619) 425-2972",
            "A/C No.                                        ": " 4587236908230087",
            "SSN                                               ": " 653-30-9562"
        }

        for key, value in customer_info.items():
            doc.add_paragraph(f"{key}: {value}")
        paragraph = doc.add_paragraph()
        # Add a subheader for Suspect infor to the document ->>
        subheader_paragraph = doc.add_paragraph("Suspect's Info")
        subheader_paragraph.style = "Heading 2"
        paragraph = doc.add_paragraph()
        #""" Addition of a checkbox where unticked box imply unavailability of suspect info"""

        # Add the customer information
        sent_val = "No suspect has been reported."
        paragraph = doc.add_paragraph()
        runner = paragraph.add_run(sent_val)
        runner.bold = True
        runner.italic = True
        suspect_info = {
            "Name                                           ": "",
            "Address                                      ": "",
            "Phone                                          ": "",
            "SSN                                               ": "",
            "Relationship with Customer ": ""
        }

        for key, value in suspect_info.items():
            doc.add_paragraph(f"{key}: {value}")
        
        doc.add_heading('Summary', level=2)
        paragraph = doc.add_paragraph()
        doc.add_paragraph(st.session_state["tmp_summary"])
        paragraph = doc.add_paragraph()
        doc.add_heading('Key Insights', level=2)
        paragraph = doc.add_paragraph()
        st.session_state.tmp_table.drop_duplicates(inplace=True)
        columns = list(st.session_state.tmp_table.columns)
        table = doc.add_table(rows=1, cols=len(columns), style="Table Grid")
        table.autofit = True
        for col in range(len(columns)):
            # set_cell_margins(table.cell(0, col), top=100, start=100, bottom=100, end=50) # set cell margin
            table.cell(0, col).text = columns[col]
        # doc.add_table(st.session_state.tmp_table.shape[0]+1, st.session_state.tmp_table.shape[1], style='Table Grid')
        
        for i, row in enumerate(st.session_state.tmp_table.itertuples()):
            table_row = table.add_row().cells # add new row to table
            for col in range(len(columns)): # iterate over each column in row and add text
                table_row[col].text = str(row[col+1]) # avoid index by adding col+1
        # save document
        # output_bytes = docx.Document.save(doc, 'output.docx')
        # st.download_button(label='Download Report', data=output_bytes, file_name='evidence.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        bio = io.BytesIO()
        doc.save(bio)
        # col_d1, col_d2 = st.columns(2)
        col_d1, col_d2 = st.tabs(["Download Report", "Download Case Package"])
        with col_d1:
        # Applying to download button -> download_button
            st.markdown("""
                <style>
                    .stButton download_button {
                        width: 100%;
                        height: 70%;
                    }
                </style>
            """, unsafe_allow_html=True)


            # combined_doc_path = os.path.join(tmp_dir, "resulting_document.docx")
            # doc.save(combined_doc_path)

            # # Create a zip file with the uploaded PDF files and the combined document
            # zip_file_name = "package_files.zip"
            # if pdf_files:
            #     st.write(file_paths)
            #     files =  [combined_doc_path]
            #     st.write(files)
                
            #     create_zip_file(files, zip_file_name)
            #     # create_zip_file(file_paths, zip_file_name)
            # else:
            #     pass
            # # Download the package
            # with open(zip_file_name, "rb") as file:
            #     st.download_button(
            #         label="Download Case Package", 
            #         data=file, 
            #         file_name=zip_file_name,
            #         disabled=st.session_state.disabled)
            
            if doc:
                st.download_button(
                    label="Download Report",
                    data=bio.getvalue(),
                    file_name="Report.docx",
                    mime="docx",
                    disabled=st.session_state.disabled
                )
        with col_d2:

            # initiating a temp file
            tmp_dir = tempfile.mkdtemp()
        
            file_paths= []

            for uploaded_file in pdf_files:
                file_pth = os.path.join(tmp_dir, uploaded_file.name)
                with open(file_pth, "wb") as file_opn:
                    file_opn.write(uploaded_file.getbuffer())
                file_paths.append(file_pth)
            
            combined_doc_path = os.path.join(tmp_dir, "resulting_document.docx")
            doc.save(combined_doc_path)

            # Create a zip file with the uploaded PDF files and the combined document
            zip_file_name = "package_files.zip"
            if file_paths:
                files =  [combined_doc_path] + file_paths
                
                create_zip_file(files, zip_file_name)
                # create_zip_file(file_paths, zip_file_name)
            else:
                pass


            
            # Download the package
            try:
                with open(zip_file_name, "rb") as file:
                    st.download_button(
                        label="Download Case Package", 
                        data=file, 
                        file_name=zip_file_name,
                        disabled=st.session_state.disabled)
            except FileNotFoundError:
                pass
                # # Cleanup: Delete the temporary directory and its contents
                # for file_path in file_paths + [combined_doc_path]:
                #     os.remove(file_path)
                # os.rmdir(temp_dir)
            
    except NameError:
        pass
        
# Adding Radio button
st.header("Make Decision")
st.markdown(
        """ <style>
                div[role="radiogroup"] >  :first-child{
                    display: none !important;
                }
            </style>
            """,
        unsafe_allow_html=True
    )
st.markdown("##### Is SAR filing required?")
selected_rad = st.radio(":blue[Please select your choice]", ["opt1","Yes", "No", "Refer for review"], horizontal=True,disabled=st.session_state.disabled)
if selected_rad == "Refer for review":
    email_regex = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    email_id = st.text_input("Enter your email ID")
    if email_id and not re.match(email_regex, email_id):
        st.error("Please enter a valid email ID")
if st.button("Submit"):
    if selected_rad in ("str_opt1"):
        st.write("")
    elif selected_rad in ("Yes"):
        st.warning("Thanks for your review, your response has been submitted")
    elif selected_rad in ("No"):
        st.success("Thanks for your review, your response has been submitted")

    else:
        st.info("Thanks for your review, Case has been assigned to the next reviewer")


# Allow the user to clear all stored conversation sessions
# if st.button("Reset Session"):
#     reset_session_state()
#     pdf_files.clear()

# Footer
st.markdown(
    """
    <style>
      #MainMenu {visibility: hidden;}
      footer {visibility: hidden;}
    </style>
    """
    , unsafe_allow_html=True)
st.markdown('<div class="footer"><p></p></div>', unsafe_allow_html=True)

hide_st_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            header {visibility: hidden;}
            </style>
            """
st.markdown(hide_st_style, unsafe_allow_html=True)


padding = 0
st.markdown(f""" <style>
    .reportview-container .main .block-container{{
        padding-top: {padding}rem;
        padding-right: {padding}rem;
        padding-left: {padding}rem;
        padding-bottom: {padding}rem;
    }} </style> """, unsafe_allow_html=True)



